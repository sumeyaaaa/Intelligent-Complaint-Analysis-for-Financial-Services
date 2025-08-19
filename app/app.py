import streamlit as st
import pandas as pd
import os
from datetime import date
from io import BytesIO
import matplotlib.pyplot as plt
import seaborn as sns
import docx
from fpdf import FPDF
import sys
import uuid
import time
from pathlib import Path
import plotly.express as px
import plotly.graph_objects as go

# --- Set a consistent project root ---
# This goes up one level from the current script's directory ('app')
project_root = Path(__file__).parent.parent
# --- Set data and evaluation log paths ---
data_path = os.path.join(project_root, "data", "clean", "complaints_clean.csv")
eval_log_path = os.path.join(project_root, "evaluation", "eval_from_chatbot.csv")

# Import RAG pipeline
src_path = os.path.join(project_root, "src")
if src_path not in sys.path:
    sys.path.append(src_path)
from RAG_pipeline_eval import RAGPipeline

# Cache RAG pipeline for performance
@st.cache_resource
def get_rag_pipeline():
    return RAGPipeline()

# Handle the case where the RAG pipeline might fail to load
try:
    rag = get_rag_pipeline()
except Exception as e:
    st.error(f"Failed to load the RAG pipeline. The app will run without the chatbot. Error: {e}")
    rag = None

# Streamlit Page Configuration & Theming
st.set_page_config(
    page_title="Intelligent Complaint Analysis",
    layout="wide",
    initial_sidebar_state="auto"
)
st.markdown("""
<style>
/* General App Style */
.st-emotion-cache-18ni7ap {
    background-color: #f0f2f6; /* Lighter background for a modern feel */
}
.stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
    font-size: 1.25rem;
    font-weight: bold;
}
/* Style for subheaders */
h3 {
    color: #0d47a1;
}
/* Custom style for metric numbers */
div[data-testid="stMetric"] label p {
    color: #0d47a1;
    font-weight: bold;
}
</style>
""", unsafe_allow_html=True)


# State Management Initialization
if 'messages' not in st.session_state:
    st.session_state.messages = []
if 'data_chatbot_messages' not in st.session_state:
    st.session_state.data_chatbot_messages = []
if 'date_range' not in st.session_state:
    st.session_state.date_range = (date(2022, 1, 1), date.today())

# Data Loading (Cached for Performance)
@st.cache_data
def load_and_preprocess_data():
    """Loads and caches the complaint data."""
    try:
        df = pd.read_csv(data_path)
        df["Date received"] = pd.to_datetime(df["Date received"])
        return df
    except FileNotFoundError:
        st.error(f"Data file not found at: {data_path}. Please check your project structure.")
        return pd.DataFrame()

# Callback function to update session state
def update_date_range():
    st.session_state.date_range = st.session_state['date_picker']

# Function to save evaluation
def save_evaluation(question, answer, rating, comment):
    # Ensure the directory exists
    os.makedirs(os.path.dirname(eval_log_path), exist_ok=True)
    
    # Check if the file exists and is not empty
    file_exists = os.path.exists(eval_log_path) and os.path.getsize(eval_log_path) > 0
    
    new_record = {
        "timestamp": [pd.Timestamp.now()],
        "question": [question],
        "answer": [answer],
        "rating": [rating],
        "comment": [comment]
    }
    new_df = pd.DataFrame(new_record)

    # Append to the CSV file
    if file_exists:
        new_df.to_csv(eval_log_path, mode='a', header=False, index=False)
    else:
        new_df.to_csv(eval_log_path, mode='w', header=True, index=False)
        
    st.success("✅ Your feedback has been saved!")

# Page Header
st.title("Intelligent Complaint Analysis 🧠📊")
st.markdown("### A comprehensive platform for exploring loan complaint data with an intelligent chatbot.")

# The Main Date Filter (placed at the top)
st.markdown("---")
col1, col2 = st.columns([1, 4])
with col1:
    date_range = st.date_input(
        "Choose a Date Range",
        st.session_state.date_range,
        min_value=date(2010, 1, 1),
        max_value=date.today(),
        key='date_picker',
        on_change=update_date_range
    )
    if isinstance(date_range, date):
        date_range = (date_range, date_range)
    st.session_state.date_range = date_range

# Load and filter data based on selected date range
df_full = load_and_preprocess_data()
if df_full.empty:
    st.stop()

start_date, end_date = st.session_state.date_range
filtered_df = df_full[
    (df_full["Date received"] >= pd.to_datetime(start_date)) &
    (df_full["Date received"] <= pd.to_datetime(end_date))
]

st.info(f"Showing analysis for **{len(filtered_df):,}** complaints from **{start_date}** to **{end_date}**.")
st.markdown("---")

# Tabs for different sections
tab1, tab2, tab3 = st.tabs(["📊 Analysis & Reports", "📈 Visualizations", "💬 Chatbot"])

with tab1:
    st.header("Analysis Report")
    if filtered_df.empty:
        st.warning("No data found for the selected date range. Please adjust the dates.")
    else:
        def create_pdf_report(df, date_range):
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt="Loan Complaint Analysis Report", ln=1, align="C")
            pdf.set_font("Arial", size=10)
            pdf.cell(200, 10, txt=f"Date Range: {date_range[0]} to {date_range[1]}", ln=1, align="C")
            pdf.ln(5)
            pdf.set_font("Arial", style='B', size=11)
            pdf.cell(200, 10, txt="Key Metrics:", ln=1)
            pdf.set_font("Arial", size=10)
            pdf.multi_cell(0, 5, f"Total Complaints: {len(df)}")
            pdf.multi_cell(0, 5, f"Companies with most complaints:\n{df['Company'].value_counts().head(5).to_string()}")
            pdf.multi_cell(0, 5, f"Top 5 products:\n{df['Product'].value_counts().head(5).to_string()}")
            
            # --- Add new metrics to the PDF report ---
            pdf.ln(5)
            pdf.set_font("Arial", style='B', size=11)
            pdf.cell(200, 10, txt="Complaint Insights:", ln=1)
            pdf.set_font("Arial", size=10)
            
            # Consumer consent
            consent_counts = df['Consumer consent provided?'].value_counts(normalize=True) * 100
            consent_text = f"Consumer Consent Provided:\n{consent_counts.round(2).to_string()} %"
            pdf.multi_cell(0, 5, consent_text)
            
            # Submission method
            submit_counts = df['Submitted via'].value_counts()
            submit_text = f"Top Submission Methods:\n{submit_counts.head(3).to_string()}"
            pdf.multi_cell(0, 5, submit_text)
            
            # Timely response
            timely_counts = df['Timely response?'].value_counts(normalize=True) * 100
            timely_text = f"Timely Response Status:\n{timely_counts.round(2).to_string()} %"
            pdf.multi_cell(0, 5, timely_text)
            
            pdf_bytes = pdf.output(dest='S').encode('latin1')
            return BytesIO(pdf_bytes)

        st.subheader("Key Metrics")
        col_m1, col_m2, col_m3 = st.columns(3)
        with col_m1:
            st.metric("Total Complaints", f"{len(filtered_df):,}")
        with col_m2:
            st.metric("Top Product", filtered_df['Product'].mode()[0])
        with col_m3:
            st.metric("Top Company", filtered_df['Company'].mode()[0])
        st.subheader("Data Table (Filtered)")
        # Limit the dataframe to the top 20 rows
        st.dataframe(filtered_df.head(20))
        if st.button("Generate & Download PDF Report"):
            pdf_bytes = create_pdf_report(filtered_df, st.session_state.date_range)
            st.download_button(
                label="Download PDF",
                data=pdf_bytes,
                file_name="analysis_report.pdf",
                mime="application/pdf",
                type="primary"
            )

        st.markdown("---")
        st.subheader("Data Chatbot")
        st.info("Ask me questions about the complaints within the selected date range.")

        # Chatbot UI for Analysis & Reports
        for message in st.session_state.data_chatbot_messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        if prompt := st.chat_input("Ask a question about the filtered data...", key="data_chatbot_input"):
            st.session_state.data_chatbot_messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    try:
                        def process_rag_query(query, df):
                            if "number of complaints" in query.lower():
                                return f"There are **{len(df):,}** complaints in the selected date range."
                            elif "top product" in query.lower():
                                top_product = df['Product'].mode()[0]
                                return f"The top product with the most complaints is **{top_product}**."
                            elif "top company" in query.lower():
                                top_company = df['Company'].mode()[0]
                                return f"The company with the most complaints is **{top_company}**."
                            elif "how many companies" in query.lower():
                                num_companies = df['Company'].nunique()
                                return f"There are **{num_companies}** unique companies in this dataset."
                            else:
                                return "I can only answer simple questions about the number of complaints, top products, and top companies."

                        response = process_rag_query(prompt, filtered_df)
                        st.markdown(response)
                        st.session_state.data_chatbot_messages.append({"role": "assistant", "content": response})
                    except Exception as e:
                        st.error(f"An error occurred: {e}")
                        st.session_state.data_chatbot_messages.append({"role": "assistant", "content": "Sorry, I couldn't analyze the data for that question."})


with tab2:
    st.header("Data Visualizations 📈")
    if filtered_df.empty:
        st.warning("No data found for the selected date range. Please adjust the dates.")
    else:
        # --- Complaints Over Time (Plotly Line Chart) ---
        st.subheader("Complaints Over Time")
        time_counts = filtered_df.set_index("Date received").resample("M").size().reset_index()
        time_counts.columns = ['Date', 'Complaint Count']
        fig_time = px.line(
            time_counts,
            x='Date',
            y='Complaint Count',
            title='Complaints Over Time (Monthly)',
            labels={'Complaint Count': 'Number of Complaints'},
            template='plotly_white'
        )
        fig_time.update_traces(marker=dict(size=8, color="#0d47a1"))
        st.plotly_chart(fig_time, use_container_width=True)

        # --- Two-column layout for the next set of plots ---
        col_v1, col_v2 = st.columns(2)
        
        with col_v1:
            # --- Complaints by Product (Plotly Bar Chart) ---
            st.subheader("Complaints by Product")
            product_counts = filtered_df['Product'].value_counts().head(10).reset_index()
            product_counts.columns = ['Product', 'Count']
            fig_product = px.bar(
                product_counts,
                x='Count',
                y='Product',
                orientation='h',
                title='Top 10 Products by Complaint Count',
                labels={'Count': 'Number of Complaints', 'Product': 'Product Type'},
                color='Count',
                color_continuous_scale=px.colors.sequential.Teal,
                template='plotly_white'
            )
            fig_product.update_layout(yaxis={'categoryorder':'total ascending'})
            st.plotly_chart(fig_product, use_container_width=True)

        with col_v2:
            # --- New Visualization: Consumer Consent Provided? (Plotly Pie Chart) ---
            st.subheader("Consumer Consent Provided?")
            consent_counts = filtered_df['Consumer consent provided?'].value_counts().reset_index()
            consent_counts.columns = ['Consent', 'Count']
            fig_consent = px.pie(
                consent_counts,
                names='Consent',
                values='Count',
                hole=0.4,
                title='Percentage of Complaints with Consumer Consent',
                template='plotly_white',
                color_discrete_sequence=px.colors.qualitative.Pastel
            )
            st.plotly_chart(fig_consent, use_container_width=True)
            
        # --- Treemap and Sunburst Chart (New, more advanced visualizations) ---
        st.markdown("---")
        st.subheader("Advanced Hierarchical Visualizations")
        col_v3, col_v4 = st.columns(2)
        
        with col_v3:
            st.subheader("Complaints by Product & Sub-product")
            # Prepare data for Treemap (Product -> Sub-product)
            treemap_data = filtered_df.groupby(['Product', 'Sub-product']).size().reset_index(name='count')
            fig_treemap = px.treemap(
                treemap_data, 
                path=['Product', 'Sub-product'], 
                values='count',
                title='Complaint Volume by Product & Sub-product',
                color='count',
                color_continuous_scale=px.colors.sequential.Viridis
            )
            st.plotly_chart(fig_treemap, use_container_width=True)
            
        with col_v4:
            st.subheader("Complaints by Company & Product")
            # Prepare data for Sunburst (Company -> Product)
            sunburst_data = filtered_df.groupby(['Company', 'Product']).size().reset_index(name='count')
            fig_sunburst = px.sunburst(
                sunburst_data.head(50), # Limit for performance
                path=['Company', 'Product'], 
                values='count',
                title='Top Companies & Their Complaint Products',
                color='count',
                color_continuous_scale=px.colors.sequential.Sunset
            )
            st.plotly_chart(fig_sunburst, use_container_width=True)

        # Display the evaluation data table
        st.markdown("---")
        st.subheader("Chatbot Evaluation Data")
        if os.path.exists(eval_log_path) and os.path.getsize(eval_log_path) > 0:
            eval_df = pd.read_csv(eval_log_path)
            eval_df['timestamp'] = pd.to_datetime(eval_df['timestamp'])
            eval_df = eval_df.sort_values(by='timestamp', ascending=False).head(20)
            st.dataframe(eval_df)
        else:
            st.info("No chatbot evaluation data has been saved yet.")

with tab3:
    st.header("Intelligent Chatbot")
    st.info("Ask me questions about the complaints.")

    if rag:
        # Chatbot UI
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])

        # This is the main section to handle user input and bot response
        if prompt := st.chat_input("Ask a question...", key="main_chatbot_input"):
            # Add user message to chat history
            st.session_state.messages.append({"role": "user", "content": prompt})
            
            # Display user message in chat message container
            with st.chat_message("user"):
                st.markdown(prompt)

            # Display assistant response in chat message container
            with st.chat_message("assistant"):
                with st.spinner("Thinking..."):
                    try:
                        # The RAG pipeline generates the answer and other info
                        generated_answer, retrieved_chunks, sentiment_summary = rag.query(prompt)
                        
                        # Construct the full response for display
                        full_response = (
                            f"**Generated Answer:**\n{generated_answer}\n\n"
                            f"--- Sentiment of Sources ---\n📊 {sentiment_summary}\n\n"
                            f"**Retrieved Narratives:**\n\n" +
                            "\n\n".join([f"**- Narrative {i+1}:** {chunk['chunk_text']}" for i, chunk in enumerate(retrieved_chunks)])
                        )
                        
                        # Display the full response
                        st.markdown(full_response)
                        
                        # Append the assistant's full response to the messages state for history
                        unique_id = str(uuid.uuid4())
                        st.session_state.messages.append({
                            "role": "assistant",
                            "content": full_response,
                            "unique_id": unique_id,
                            "prompt": prompt,
                            "answer": generated_answer,
                            "rating": None,
                            "comment": None
                        })
                        st.rerun()

                    except Exception as e:
                        st.error(f"An error occurred with the RAG pipeline: {e}")
                        st.session_state.messages.append({"role": "assistant", "content": "I apologize, I could not process your request at this time."})
                        st.rerun()

        # Display Evaluation Form for the Last Assistant Message
        if st.session_state.messages and st.session_state.messages[-1]["role"] == "assistant":
            last_message = st.session_state.messages[-1]
            
            if not last_message.get("rating"): # Use .get() to safely check for the key
                with st.form("evaluation_form", clear_on_submit=False):
                    st.subheader("Rate this response")
                    
                    rating = st.slider(
                        "Rating (1-5)",
                        min_value=1,
                        max_value=5,
                        step=1,
                        key=f"rating_slider_{last_message['unique_id']}"
                    )
                    
                    comment = st.text_area(
                        "Why did you give this rating?",
                        key=f"comment_area_{last_message['unique_id']}"
                    )
                    
                    submit_button = st.form_submit_button("Submit Rating")
                    
                    if submit_button:
                        last_message["rating"] = rating
                        last_message["comment"] = comment
                        
                        save_evaluation(
                            last_message["prompt"],
                            last_message["answer"],
                            rating,
                            comment
                        )
                        
                        st.success("✅ Feedback submitted!")
                        time.sleep(1)
                        st.rerun()

        # Download Chat History
        if st.session_state.messages:
            def create_word_doc(history):
                doc = docx.Document()
                doc.add_heading("Complaint Chatbot Report", level=1)
                for entry in history:
                    doc.add_paragraph(f"{entry['role'].capitalize()}: {entry['content']}")
                    if 'rating' in entry and entry['rating'] is not None:
                        doc.add_paragraph(f"Rating: {entry['rating']}/5")
                    if 'comment' in entry and entry['comment']:
                        doc.add_paragraph(f"Comment: {entry['comment']}")
                bio = BytesIO()
                doc.save(bio)
                return bio.getvalue()

            doc_file = create_word_doc(st.session_state.messages)
            st.download_button(
                label="Download Chat History (docx)",
                data=doc_file,
                file_name="chatbot_history.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("The RAG chatbot is unavailable due to a loading error.")
        st.info("You can still use the 'Analysis & Reports' and 'Visualizations' tabs.")