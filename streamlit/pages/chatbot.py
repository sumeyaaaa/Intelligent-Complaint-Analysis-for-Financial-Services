import streamlit as st
import pandas as pd
import os
import sys
import io
import docx

# --- Add src directory to path ---
project_root = os.path.abspath(os.path.join(os.getcwd(), ".."))
src_path = os.path.join(project_root, "src")
if src_path not in sys.path:
    sys.path.append(src_path)

from rag_pipeline import RAGPipeline

# --- Streamlit Setup ---
st.set_page_config(page_title="RAG Chatbot", layout="centered")
st.title("💬 Complaint Chatbot")
st.markdown("Ask questions about customer complaints.")

# --- Load RAG pipeline (and cache it for performance) ---
@st.cache_resource
def get_rag_pipeline():
    return RAGPipeline(top_k=5)

rag = get_rag_pipeline()

# --- Session state setup for chat history ---
if "messages" not in st.session_state:
    st.session_state.messages = []
if "eval_df" not in st.session_state:
    if os.path.exists("evaluation_log.csv"):
        st.session_state.eval_df = pd.read_csv("evaluation_log.csv")
    else:
        st.session_state.eval_df = pd.DataFrame(columns=["question", "answer", "source_1", "source_2", "score", "comment"])

# --- Function to save evaluation ---
def save_evaluation(question, answer, chunks, score, comment):
    record = {
        "question": question,
        "answer": answer,
        "source_1": chunks[0] if len(chunks) > 0 else "",
        "source_2": chunks[1] if len(chunks) > 1 else "",
        "score": score,
        "comment": comment
    }
    st.session_state.eval_df = pd.concat([st.session_state.eval_df, pd.DataFrame([record])], ignore_index=True)
    st.session_state.eval_df.to_csv("evaluation_log.csv", index=False)
    st.success("✅ Evaluation saved.")

# --- Display chat messages from history on app rerun ---
for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.markdown(message["content"])

        if "chunks" in message and "rating_submitted" not in message:
            with st.expander("📄 Supporting Sources"):
                for i, chunk in enumerate(message["chunks"]):
                    st.write(f"**Chunk {i+1}**")
                    st.markdown(chunk)

            # --- Evaluation block ---
            with st.container():
                st.subheader("📝 Evaluate This Answer")
                col1, col2 = st.columns([1, 2])
                with col1:
                    score = st.radio("Rate (1-5)", [1, 2, 3, 4, 5], key=f"score_{message['timestamp']}", horizontal=True)
                with col2:
                    comment = st.text_input("Feedback", placeholder="Leave a comment...", key=f"comment_{message['timestamp']}")
                
                if st.button("✅ Submit Evaluation", key=f"save_{message['timestamp']}"):
                    save_evaluation(message["question"], message["content"], message["chunks"], score, comment)
                    message["rating_submitted"] = True
                    st.rerun()

# --- React to user input ---
if prompt := st.chat_input("Ask a question about complaints..."):
    # Add user message to chat history
    st.session_state.messages.append({"role": "user", "content": prompt})
    
    # Display user message in chat message container
    with st.chat_message("user"):
        st.markdown(prompt)

    # Get RAG response
    with st.spinner("Thinking..."):
        answer, chunks = rag.generate_answer(prompt)

    # Add assistant response to chat history
    st.session_state.messages.append({
        "role": "assistant",
        "content": answer,
        "chunks": chunks,
        "question": prompt,
        "timestamp": datetime.now().isoformat()
    })
    st.rerun()

# ---
st.markdown("---")
# --- Download button ---
if st.session_state.messages:
    # Function to create a Word document from chat history
    def create_word_doc(history):
        doc = docx.Document()
        doc.add_heading("Complaint Chatbot Report", level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("---")

        for entry in history:
            if entry["role"] == "user":
                doc.add_heading("Question:", level=3)
                doc.add_paragraph(entry["content"])
            elif entry["role"] == "assistant":
                doc.add_heading("Answer:", level=3)
                doc.add_paragraph(entry["content"])
                if "chunks" in entry:
                    doc.add_paragraph("\nSupporting Sources:")
                    for i, chunk in enumerate(entry["chunks"]):
                        doc.add_paragraph(f"Chunk {i+1}: {chunk[:200]}...")  # Add a snippet of the chunk

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue()

    # Create the download button
    doc_file = create_word_doc(st.session_state.messages)
    st.download_button(
        label="Download Chat Report (Word)",
        data=doc_file,
        file_name="chatbot_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )